"""
Ynov'iT Presales Pipeline — Générateur CDC Word

Transforme la sortie structurée de l'agent CDC en document
Word (.docx) formaté et professionnel.
"""

import logging
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger("presales.generators.docx")

# Couleurs Ynov'iT
NAVY = RGBColor(0x1A, 0x2B, 0x4A)
TEAL = RGBColor(0x3E, 0xC9, 0xA7)
GRAY = RGBColor(0x4A, 0x5C, 0x7A)


def generate_cdc_docx(cdc_data: dict, output_path: Path, societe: dict = None):
    """
    Génère le cahier des charges en fichier Word.

    Args:
        cdc_data: Sortie de l'agent CDC (titre, sections)
        output_path: Chemin du fichier .docx à créer
        societe: Infos société pour la page de garde
    """
    doc = Document()

    # ── Styles ─────────────────────────────────────────────
    _setup_styles(doc)

    # ── Page de garde ──────────────────────────────────────
    _add_cover_page(doc, cdc_data, societe)
    doc.add_page_break()

    # ── Sommaire (placeholder) ─────────────────────────────
    doc.add_paragraph("SOMMAIRE", style="Heading 1")
    doc.add_paragraph(
        "(Le sommaire sera généré automatiquement à l'ouverture dans Word : "
        "clic droit sur ce texte → Mettre à jour les champs)",
        style="Normal"
    )
    doc.add_page_break()

    # ── Sections ───────────────────────────────────────────
    sections = cdc_data.get("sections", [])
    for section in sections:
        _add_section(doc, section)

    # ── Sauvegarde ─────────────────────────────────────────
    doc.save(str(output_path))
    logger.info(f"CDC Word généré : {output_path} ({len(sections)} sections)")


def _setup_styles(doc):
    """Configure les styles du document."""
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level, size in [(1, 18), (2, 14), (3, 12)]:
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.name = "Arial"
        h_style.font.size = Pt(size)
        h_style.font.color.rgb = NAVY
        h_style.font.bold = True
        h_style.paragraph_format.space_before = Pt(18)
        h_style.paragraph_format.space_after = Pt(8)


def _add_cover_page(doc, cdc_data, societe):
    """Ajoute la page de garde."""
    # Espace vertical
    for _ in range(6):
        doc.add_paragraph()

    # Logo / Nom société
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("YNOV'IT SERVICES")
    run.font.size = Pt(14)
    run.font.color.rgb = TEAL
    run.font.bold = True

    doc.add_paragraph()

    # Titre
    titre = cdc_data.get("titre", "Cahier des charges")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(titre)
    run.font.size = Pt(24)
    run.font.color.rgb = NAVY
    run.font.bold = True

    doc.add_paragraph()

    # Infos société
    if societe:
        nom = societe.get("raison_sociale", "")
        secteur = societe.get("secteur_activite", "")
        if nom:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(nom)
            run.font.size = Pt(16)
            run.font.color.rgb = GRAY
        if secteur:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(secteur)
            run.font.size = Pt(12)
            run.font.color.rgb = GRAY

    # Date et version
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version = cdc_data.get("version", "1.0")
    run = p.add_run(
        f"Version {version} — {cdc_data.get('date', date.today().strftime('%d/%m/%Y'))}"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Document confidentiel")
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    run.font.italic = True


def _add_section(doc, section, level=1):
    """Ajoute une section et ses sous-sections récursivement."""
    numero = section.get("numero", "")
    titre = section.get("titre", "")
    contenu = section.get("contenu", "")

    # Titre de section
    heading_text = f"{numero}. {titre}" if numero else titre
    doc.add_paragraph(heading_text, style=f"Heading {min(level, 3)}")

    # Contenu
    if contenu:
        doc.add_paragraph(contenu, style="Normal")

    # Sous-sections
    for sub in section.get("sous_sections", []):
        _add_section(doc, sub, level=level + 1)
